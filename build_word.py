from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor
from docx.text.run import Run


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "GOAI_四页提交稿_Math_Structurer.docx"
VIS = ROOT / "artifacts" / "visuals"

NAVY = "0D2538"
BLUE = "245F83"
ORANGE = "E77859"
PALE = "EDF3F4"
PALE_ORANGE = "FBECE7"
GRAY = "66727D"
LIGHT = "F7F8F7"
WHITE = "FFFFFF"


# Author-year references are real external hyperlinks in the generated Word file.
# Keep the longest aliases first so a shorter label cannot split a longer citation.
CITATION_LINKS = sorted(
    [
        ("Güngör, arXiv:1901.01543v11, 2025", "https://arxiv.org/abs/1901.01543"),
        ("Mialon et al., 2023（arXiv v2, 2024）", "https://arxiv.org/abs/2307.05432"),
        ("Murugan & Palanivel, 2021", "https://doi.org/10.1007/s00010-020-00739-w"),
        ("Deng, Hani & Ma, 2025", "https://arxiv.org/abs/2408.07818"),
        ("Brandstetter et al., 2022", "https://proceedings.mlr.press/v162/brandstetter22a.html"),
        ("Wang et al., 2021", "https://doi.org/10.1021/acscatal.1c01504"),
        ("Wang et al. (2021)", "https://doi.org/10.1021/acscatal.1c01504"),
        ("Pd₁/Fe₃O₄ 2018", "https://doi.org/10.1002/cctc.201800362"),
        ("Cu@Na-Beta 2020", "https://doi.org/10.1016/j.chempr.2020.07.001"),
        ("alphaXiv 文档", "https://www.alphaxiv.org/docs/mcp"),
        ("Catalysis-Hub（Winther et al., 2019）", "https://doi.org/10.1038/s41597-019-0081-y"),
        ("OC20（Chanussot et al., 2021）", "https://doi.org/10.1021/acscatal.0c04525"),
        ("Materials Project API", "https://docs.materialsproject.org/downloading-data/using-the-api/getting-started"),
        ("Materials Project OPTIMADE", "https://optimade.materialsproject.org/v1/structures?filter=id%3D%22mp-19306%22&page_limit=1&response_fields=id,chemical_formula_reduced,lattice_vectors,cartesian_site_positions,species_at_sites"),
        ("Crossref REST API", "https://www.crossref.org/documentation/retrieve-metadata/rest-api/"),
        ("OpenAlex Works API", "https://help.openalex.org/data/works/attributes.md"),
        ("alphaXiv MCP", "https://alphaxiv.org/docs/mcp"),
        ("arXiv API", "https://info.arxiv.org/help/api/index.html"),
        ("Kozyra, 2021", "https://doi.org/10.20429/tag.2021.080108"),
        ("Murugan–Palanivel 2021", "https://doi.org/10.1007/s00010-020-00739-w"),
        ("Deng–Hani–Ma 2025", "https://arxiv.org/abs/2408.07818"),
        ("Brandstetter 2022", "https://proceedings.mlr.press/v162/brandstetter22a.html"),
        ("Mialon 2023／v2 2024", "https://arxiv.org/abs/2307.05432"),
        ("Güngör v11 2025", "https://arxiv.org/abs/1901.01543"),
        ("Kozyra 2021", "https://doi.org/10.20429/tag.2021.080108"),
        ("Wang 2021", "https://doi.org/10.1021/acscatal.1c01504"),
    ],
    key=lambda item: len(item[0]),
    reverse=True,
)


def font_run(run, size=8.0, bold=False, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, size=8.0, color=BLUE, bold=False, font="Microsoft YaHei"):
    """Append a real external Word hyperlink and return its Run wrapper."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run_element = OxmlElement("w:r")
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    run = Run(run_element, paragraph)
    run.text = text
    font_run(run, size, bold, color, font)
    run.font.underline = True
    return run


def add_text_with_links(paragraph, text, size=8.0, color=None, bold=False, font="Microsoft YaHei"):
    """Append text while turning every registered author-year label into a hyperlink."""
    cursor = 0
    while cursor < len(text):
        matches = []
        for label, url in CITATION_LINKS:
            index = text.find(label, cursor)
            if index >= 0:
                matches.append((index, -len(label), label, url))
        if not matches:
            run = paragraph.add_run(text[cursor:])
            font_run(run, size, bold, color, font)
            break
        index, _, label, url = min(matches)
        if index > cursor:
            run = paragraph.add_run(text[cursor:index])
            font_run(run, size, bold, color, font)
        add_hyperlink(paragraph, label, url, size, BLUE, bold, font)
        cursor = index + len(label)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=32, start=55, bottom=32, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def keep_next(paragraph):
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(7.9)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(1.5)

    for name, size, color in (
        ("Title", 16.5, NAVY),
        ("Heading 1", 12.5, NAVY),
        ("Heading 2", 9.8, BLUE),
        ("Heading 3", 8.7, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(2.0)
        style.paragraph_format.space_after = Pt(1.4)
        style.paragraph_format.keep_with_next = True

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.00)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.35)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("GOAI 开放探索  •  ")
    font_run(r, 6.8, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)


def add_h(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_next(p)
    return p


def add_p(doc, text, size=7.9, color=None, bold_lead=None, after=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        font_run(r, size, True, color)
        add_text_with_links(p, text[len(bold_lead):], size, color)
    else:
        add_text_with_links(p, text, size, color)
    return p


def add_formula(doc, text, size=9.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0.8)
    p.paragraph_format.space_after = Pt(1.4)
    r = p.add_run(text)
    font_run(r, size, False, NAVY, "Cambria Math")
    return p


def add_table(doc, headers, rows, widths, font_size=7.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    cant_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        font_run(r, font_size, True, WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        cant_split(table.rows[-1])
        for index, value in enumerate(values):
            cell = cells[index]
            cell.width = Cm(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            if row_index % 2:
                shade(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text_with_links(p, str(value), font_size, bold=index == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    return table


def add_picture(doc, name, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0.7)
    p.paragraph_format.space_after = Pt(1.0)
    p.add_run().add_picture(str(VIS / name), width=Cm(width_cm))


def add_meta(doc, fields, citations, visual, human, budget):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    entries = [
        ("页面字段：", fields),
        ("  文献依据：", citations),
        ("  建议视觉：", visual),
        ("  人工确认：", human),
        ("  内容预算：", budget),
    ]
    for lead, body in entries:
        r = p.add_run(lead)
        font_run(r, 6.45, True, BLUE)
        add_text_with_links(p, body, 6.45, GRAY)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    setup(doc)
    props = doc.core_properties
    props.title = "Math Structurer — GOAI 开放探索四页问题定义"
    props.subject = "自然语言反应目标的证据、能量与几何结构化，以及特定基算子迭代"
    props.author = "GOAI 参赛团队"
    props.keywords = "GOAI, Math Structurer, 催化剂, 反应能量, 空间几何, 特定基算子, 迭代, 文献核验"

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Math Structurer")
    font_run(r, 16.5, True, NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    r = subtitle.add_run("Convincing, reusable target-matching skills for AI research agents.")
    font_run(r, 8.4, True, BLUE)

    add_table(
        doc,
        ["四页共用信息", "冻结表述"],
        [
            ("一句话问题", "科研代理能否把自然语言反应目标转换为反应物、产物、带来源反应能量、催化剂空间几何及文献记录，并利用空值、冲突和几何缺口改变下一轮检索，而不把候选记录冒充发现？"),
            ("研究边界", "首版只运行 CO2gas+H2gas -- CH3CH2OHgas @best 和独立的特定基算子调试；@best 只调用固定排序式。温度、压力、候选范围、观测指标和基线为空时仍允许查询。"),
            ("当前成熟度", "交互原型 v0.6；Windows 启动、浏览器与本地识别通过。2026-08-18 又完成 Materials Project 认证查询和 alphaXiv fullText 抽取；Linux／macOS 未实测。当前无可比反应能量、活性位几何、性能结论或催化发现。"),
        ],
        [2.7, 15.6],
        6.65,
    )

    add_h(doc, "一、问题与证据", 1)
    add_h(doc, "1.1 真实问题或需求", 2)
    add_p(doc, "化学实验中的一句自然语言目标常把反应、能量口径、催化剂记录、空间结构和排序依据混在一起。以 CO₂ 与 H₂ 生成乙醇为例，Wang 等讨论 Na–Fe@C、K–CuZnAl、界面及组分邻近方式，说明候选必须携带来源和结构上下文，却不支持由题名或元数据直接推出跨论文性能次序（Wang et al., 2021）。Pd₁/Fe₃O₄ 2018 与 Cu@Na-Beta 2020 只能先作为可追溯检索候选。", 7.6)
    add_p(doc, "数学研究也常呈现“有限结构生成候选—验证器给出反例—下一轮改变表示”。半迭代是在指定 X 上寻找 g:X→X 使 g∘g=f，其存在性依赖函数图或受限连续函数类的内部结构（Kozyra, 2021；Murugan & Palanivel, 2021）。李对称可压缩微分方程变量并支持机器学习表征；本地 MATHs/math 3624 number theory 生成器则产生轨道后检查素性、同余和闭合（Güngör, arXiv:1901.01543v11, 2025；Brandstetter et al., 2022；Mialon et al., 2023（arXiv v2, 2024））。共同点只是特定基算子＋迭代的工作流，并非已知统一代数。", 7.45)
    add_h(doc, "1.2 为什么尚未被结构化", 2)
    add_p(doc, "结构化瓶颈首先是定义不一致：反应能量可能指电子能差、焓变、自由能变、势垒或吸附能；缺少数值、单位、方法、参考态和来源时，任何数字都不可复查。催化剂又可能是分子、团簇、表面、载体或界面，二维连线、三维坐标和对称信息必须交给不同插件，并标明坐标来源与“体相／支撑体／表面／活性位”范围。", 7.5)
    add_p(doc, "纯数学侧也没有可直接套用的万能函数生成器。Deng、Hani 与 Ma 从硬球动力学长时导出 Boltzmann 方程时，用时间分层、累积量、部分展开、分子图和切割算法控制误差（Deng, Hani & Ma, 2025，pp. 9–12）。这份 192 页推导只作复杂度例证：专业参数、函数空间、结构表示与误差口径必须一起嵌入，不能由语言模型一次补全。", 7.5)
    add_h(doc, "1.3 研究价值与合适切片", 2)
    add_p(doc, "演示只建立可审计的“反应—证据—几何”记录：反应物／产物来自输入，能量为 null，五条候选保留来源；当前三维对象是 Materials Project OPTIMADE 的 mp-19306 Fe₃O₄ 体相支撑体坐标，范围 support-only，不是 Pd 活性位。", 7.3, after=0.5)
    add_formula(doc, "S(c)=50mRP+25mabs+15mgeom+10menergy　（四项为 0／1，只作稳定检索排序）", 8.0)
    add_picture(doc, "page1_demo_slice.png", 16.4)
    add_meta(
        doc,
        "反应物；产物；能量值／单位／定义／方法／来源／状态；候选记录；固定排序特征；几何来源与对象范围。",
        "Wang 2021；Kozyra 2021；Murugan–Palanivel 2021；Güngör v11 2025；Brandstetter 2022；Mialon 2023／v2 2024；Deng–Hani–Ma 2025；Materials Project OPTIMADE。",
        "v0.6 演示截图：反应物—产物—能量空值—固定排序边界—相连空间。",
        "逐项打开候选链接；未知能量保持空值；mp-19306 未写成 Pd 活性位；排序未写成性能结论。",
        "约 850–950 字＋1 张实测截图。",
    )

    page_break(doc)

    add_h(doc, "二、环境接口", 1)
    add_h(doc, "2.1 固定规则", 2)
    add_p(doc, "环境固定接收原始查询 q，并输出反应物 R、产物 P、能量记录 E、催化剂候选 C、空间几何 G 与来源集合 L。", 7.7, after=0.6)
    add_formula(doc, "q ↦ (R,P,E,C,G,L),　E=(v,u,ρ,m,s,σ)", 9.1)
    add_p(doc, "E 中 v 为数值、u 为单位、ρ 为能量定义、m 为实验或计算方法、s 为来源、σ 为核验状态；任何必要子字段缺失时 v=null，界面显示“—”。候选至少保存名称、记录标识、网址、证据层级和检索时间；几何至少保存节点、边、坐标来源及“体相／支撑体／表面／活性位”范围。", 7.55)
    add_p(doc, "@best 只调用公开的 50／25／15／10 固定排序式，不新增科学主张。温度、压力、候选范围、观测指标和基线是可选可能性，缺失不阻断检索。各数据源职责固定如下。", 7.4, after=0.6)
    add_table(
        doc,
        ["数据源", "可读取对象", "硬边界"],
        [
            ("Catalysis-Hub（Winther et al., 2019）", "反应物／产物、带类型 DFT 能量、部分原子结构", "模式内省实测 200；无密钥记录查询 401；跨论文不可直接排序"),
            ("OC20（Chanussot et al., 2021）／Materials Project API", "初始／弛豫结构／认证体相查询", "MP 实测 200；mp-19306 解析为 mp-aaaabcoo；体相不是活性位"),
            ("Crossref REST API／OpenAlex Works API", "DOI、题名、作者、日期、落地页与开放状态", "只作身份核验和跳转，不给反应能或几何"),
            ("alphaXiv MCP／arXiv API", "认证原文抽取／规范元数据回退", "无认证 401；MCP 与 fullText=true 实测 200；默认 AI 报告仍非原文"),
        ],
        [4.3, 7.2, 6.8],
        5.85,
    )
    add_h(doc, "2.2 观察/行动/反馈", 2)
    add_table(
        doc,
        ["接口", "落地字段或操作", "反馈如何改变下一步"],
        [
            ("观察", "原查询；R/P；能量六元组；候选网址与证据层级；几何来源；剩余预算", "确定当前可核验字段与空值"),
            ("行动", "选择 DOI／alphaXiv／公共数据库；改检索式；读摘要或正文；调用能量或几何插件；执行固定排序", "改变下一来源、字段或几何对象范围"),
            ("反馈", "已核验／仅元数据／缺失／冲突／接口不可用／超时；同时返回原始字段和位置", "能量缺失→授权记录接口；support-only→检索活性位；冲突→并列保留"),
        ],
        [2.2, 8.9, 7.2],
        6.55,
    )
    add_p(doc, "下一轮可改变连接器、查询词、目标证据层级或几何对象范围；固定 50／25／15／10 排序式不由 Agent 调权。至少一次反馈必须造成可见改变，模型推测不得写入已核验字段。", 7.45)
    add_h(doc, "2.3 记录与预算", 2)
    add_p(doc, "日志保存查询版本、连接器、记录标识、网址、访问时间、字段原文、能量空值原因、几何来源、固定排序特征、模型／插件版本、反馈与下一动作理由。本机 Qwen3-8B 检查点每次至多用于一次受限识别；Codex 或 DeepSeek 仅在用户显式选择时二选一、最多一次；文献与数据库接口合计八次，二维和三维输出各一幅，探索六轮。后端不可用必须明示。", 7.45)
    add_meta(
        doc,
        "固定类型；观察字段；动作枚举；反馈状态；空值语义；固定排序式；来源日志；模型、接口、几何与轮数预算。",
        "Wang 2021；Catalysis-Hub（Winther et al., 2019）；OC20（Chanussot et al., 2021）；Materials Project API；Crossref REST API；OpenAlex Works API；alphaXiv MCP；arXiv API。",
        "q→(R,P,E,C,G,L) 的类型接口与观察—行动—反馈表。",
        "可选字段为空不阻断；模型不是验证器；公共记录与科学结论分层；平台状态如实标注。",
        "约 650–750 字＋1 张接口表。",
    )

    page_break(doc)

    add_h(doc, "三、发现信号与参照", 1)
    add_h(doc, "3.1 什么算发现", 2)
    add_p(doc, "材料名称或排序首位不算发现。发现信号必须在运行前写明，并允许被来源冲突、条件扰动或专业验证器否证。", 7.65, after=0.6)
    add_table(
        doc,
        ["信号", "可操作定义", "它改变什么认识"],
        [
            ("正向发现", "同一反应定义与可比能量口径下，至少两个独立来源，或一项原始研究加专业数据库，支持同一几何—能量关系；几何对象范围明确；改预注册检索词后仍出现", "得到待专业复核的条件化结构关系；support-only 不算活性位证据"),
            ("稳定负结果", "八次预注册接口调用均没有能量—几何联合字段", "只说明当前数据覆盖不足，不说明领域不存在"),
            ("异常／反例", "同一候选的能量符号、单位、定义、参考态或结构标识在来源间冲突", "停止合并，保留两侧原文并缩小口径"),
            ("问题修正", "把“性能最佳”改为“固定检索分排序”，或先区分自由能、电子能和势垒", "下一轮改连接器、检索词或几何对象范围"),
        ],
        [2.25, 10.1, 5.95],
        6.15,
    )
    add_p(doc, "当前演示最多取得接口、空值、冲突或问题修正信号，尚未构成催化剂发现。", 7.4)
    add_h(doc, "3.2 平凡解/随机/无干预", 2)
    add_formula(doc, "𝒟q={dᵢ : match(dᵢ,R,P)>0 且 source(dᵢ)≠∅},　cᵢ=record(dᵢ)", 8.6)
    add_p(doc, "候选之所以出现，是公共记录与反应物／产物语义匹配，而不是系统证明了性能。@best 只对 𝒟q 应用固定 S(c)，不会生成缺失能量或提升证据层级。验证依次检查能量六元组，检查几何记录标识、坐标来源、节点／边与对象范围；mp-19306 只验证 Fe₃O₄ 体相支撑体，不验证 Pd 活性位。文献另分 DOI、元数据、摘要和正文层级。", 7.25)
    add_table(
        doc,
        ["参照", "同预算操作", "比较指标"],
        [
            ("平凡", "直接让语言模型给材料名", "无来源候选数、虚构能量数、边界措辞"),
            ("随机", "随机排列同一批公共记录", "与固定 S(c) 的次序和首位证据完整度比较"),
            ("无干预", "单次关键词检索，不读来源层级", "可点击来源数、空值是否保留"),
            ("非平凡", "确定性合并 DOI／公共元数据并应用固定 S(c)", "重复率、来源覆盖、能量与几何可追溯率"),
        ],
        [2.3, 8.3, 7.7],
        6.25,
    )
    add_h(doc, "3.3 最低成功与失败标准", 2)
    add_p(doc, "环境技术通过：① DOI 经 Crossref／OpenAlex 核验，alphaXiv 已在认证且 fullText=true 下完成 arXiv:2408.07818 抽取，只保存长度和哈希；② 本地验证器通过反应物／产物、能量 null、固定 50／25／15／10 排序和 mp-19306／support-only；③ 人工终审链接、公式、字段和结论边界。认证通过不等于科学发现。", 7.0, bold_lead="环境技术通过：", after=0.7)
    add_p(doc, "科学发现通过：结果须超过非平凡基线，形成可复现的几何—能量关系，或跨预注册查询稳定复现的数据覆盖负结果，并经专业工具或领域专家核验。当前科学闸门未通过。", 7.3, bold_lead="科学发现通过：", after=0.7)
    add_p(doc, "直接失败：空能量写成零；元数据冒充实验／计算结论；首位记录冒充性能最优；数据库候选冒充发现；支撑体坐标冒充活性位；可选条件缺失便停止；接口不可用写成研究不存在；模型无来源输出进入正式字段；反馈不改变行动。若条件化试跑仍不优于确定性基线，应承认当前切片不成立。", 7.15)
    add_meta(
        doc,
        "四类发现信号；能量／几何可比条件；四类参照；技术闸门；科学闸门；失败标准。",
        "Wang 2021；Crossref REST API；OpenAlex Works API；alphaXiv MCP；arXiv API；本地浏览器与结构检查收据。",
        "四类发现信号表、四类参照表和文献—本地—人工三级闸门。",
        "ACS 主页面本轮未读；当前无统一能量数值、性能结论或催化发现。",
        "约 750–850 字＋2 张紧凑表。",
    )

    page_break(doc)

    add_h(doc, "四、最小验证计划", 1)
    add_h(doc, "4.1 一次试跑怎么做", 2)
    add_p(doc, "解压后在项目根目录只执行一行：", 7.9)
    code = doc.add_paragraph()
    code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code.paragraph_format.space_after = Pt(2.0)
    r = code.add_run("python3 run.py")
    font_run(r, 10.2, True, NAVY, "Consolas")
    add_p(doc, "浏览器打开终端给出的本地地址，保留 CO2gas+H2gas -- CH3CH2OHgas @best 并点击一次“运行”。本次只验五项；离线结构检查为 python3 run.py --check。Windows v0.6 已实测，Linux／macOS 使用同一入口但仍须实机复跑。", 7.45)
    add_table(
        doc,
        ["一次试跑检查", "通过表现", "失败时保留"],
        [
            ("反应物／产物", "分别显示 CO₂、H₂ 与 C₂H₅OH", "原始输入与解析字段"),
            ("反应能量", "未知时显示“—”，内部 v=null", "空值原因"),
            ("候选来源", "五条可点击记录及证据层级", "接口状态与原字段"),
            ("@best", "严格按 50／25／15／10 固定分只改变次序", "四项特征与并列记录"),
            ("二维／三维", "mp-19306 Fe₃O₄；support-only", "不得改写为 Pd 活性位"),
        ],
        [3.2, 8.4, 6.7],
        6.35,
    )
    add_h(doc, "4.2 主要风险与失败路径", 2)
    add_p(doc, "首要风险是用可信语气覆盖空值或对象范围。能量缺定义或来源时只显示“—”；元数据不生成活性或机理；首位记录不改写成性能最佳；mp-19306 只显示为 Fe₃O₄ 体相支撑体，不补画 Pd 活性位；接口超时与来源矛盾保留原状态。模型生成的 DOI、数值或结构标识，未打开原始页面前一律待核验。", 7.25)
    add_p(doc, "数学面板独立调试特定基算子与迭代；数值点上一致不等于函数恒等，含 sorry 或额外公理的 Lean 文件不等于完成证明。该算子不处理反应能量或催化剂坐标，也不代表统一代数。", 7.3)
    add_formula(doc, "B: ℂ×ℂ^×→ℂ,　B(u,v)=eᵘ−Log(v),　E(x)=B(1,B(B(1,x),1))", 8.7)
    add_p(doc, "约定主值复对数且每个 B 节点的第二参数非零；若测试点给出反例，界面保留“不匹配”，不得改写为成功。", 7.25)
    add_h(doc, "4.3 复现与开源计划", 2)
    add_p(doc, "开源包保存冻结查询、R/P JSON、能量六元组与空值原因、候选标识、固定排序特征、mp-19306 来源及 support-only 范围、事件日志、链接、测试收据和源码哈希，不收录论文 PDF。复现画面必须显示 q↦(R,P,E,C,G,L)、函数空间 B:ℂ×ℂ^×→ℂ、复合树、反例和 Lean 状态。发布前在 Ubuntu 与 macOS 运行 python3 run.py --check；未通过的平台继续标为“未验证”。", 7.2)
    add_picture(doc, "page4_function_space_screenshot.png", 9.2)
    add_meta(
        doc,
        "单行命令；冻结输入；五项输出；空值与风险状态；平台状态；源码／日志哈希；开源边界。",
        "Wang 2021；Kozyra 2021；本轮 DOI／alphaXiv、浏览器、结构检查与 Lean 收据。",
        "v0.6 实测截图：函数空间、特定基算子复合树、反例与 Lean 状态；化学截图另示 support-only 边界。",
        "逐项打开来源；核对能量空值、固定排序式与 mp-19306 对象范围；不把设计适配写成跨平台实测。",
        "约 600–700 字＋1 行命令＋1 表＋1 张实测截图。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
